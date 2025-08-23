import PyPDF2
import pdfplumber
from docx import Document
from fastapi import UploadFile, HTTPException
import io
import re
from typing import Optional, Dict
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """Service for parsing various document formats and extracting text"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.doc', '.docx']
        logger.info("DocumentParser initialized")
    
    async def parse_document(self, file: UploadFile) -> str:
        """
        Parse uploaded document and extract text content
        
        Args:
            file: Uploaded file object
            
        Returns:
            Extracted text content
            
        Raises:
            HTTPException: If file format is unsupported or parsing fails
        """
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
            
        filename = file.filename.lower()
        logger.info(f"Parsing document: {file.filename}")
        
        try:
            # Read file content
            content = await file.read()
            logger.info(f"File size: {len(content)} bytes")
            
            if filename.endswith('.pdf'):
                text = self._parse_pdf(content)
            elif filename.endswith('.docx'):
                text = self._parse_docx(content)
            elif filename.endswith('.doc'):
                # For .doc files, try docx parser (some work)
                text = self._parse_docx(content)
            else:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Unsupported file format. Supported: {', '.join(self.supported_formats)}"
                )
            
            logger.info(f"Successfully extracted {len(text)} characters")
            return text
                
        except Exception as e:
            logger.error(f"Error parsing document {file.filename}: {str(e)}")
            raise HTTPException(
                status_code=500, 
                detail=f"Error parsing document: {str(e)}"
            )
        finally:
            # Reset file pointer for potential reuse
            await file.seek(0)
    
    def _parse_pdf(self, content: bytes) -> str:
        """Parse PDF content using multiple methods for better reliability"""
        
        # Method 1: Try pdfplumber (better for complex layouts)
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text = ""
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                if text.strip():
                    logger.info("PDF parsed successfully with pdfplumber")
                    return self._clean_text(text)
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
        
        # Method 2: Fallback to PyPDF2
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            if text.strip():
                logger.info("PDF parsed successfully with PyPDF2")
                return self._clean_text(text)
        except Exception as e:
            logger.warning(f"PyPDF2 failed: {e}")
        
        raise Exception("Unable to extract text from PDF using available methods")
    
    def _parse_docx(self, content: bytes) -> str:
        """Parse DOCX content"""
        try:
            doc = Document(io.BytesIO(content))
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise Exception("No text content found in document")
            
            logger.info("DOCX parsed successfully")
            return self._clean_text(text)
            
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with processing
        text = re.sub(r'[^\w\s\-.,@():/\n]', '', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract basic contact information from resume text"""
        contact_info = {
            "email": None,
            "phone": None,
            "linkedin": None
        }
        
        # Email regex
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info["email"] = email_match.group()
        
        # Phone regex (various formats)
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?(\d{3})\)?[-.\s]?(\d{3})[-.\s]?(\d{4})'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info["phone"] = phone_match.group()
        
        # LinkedIn profile
        linkedin_pattern = r'linkedin\.com/in/[\w\-]+'
        linkedin_match = re.search(linkedin_pattern, text.lower())
        if linkedin_match:
            contact_info["linkedin"] = linkedin_match.group()
        
        return contact_info
    
    def validate_file_size(self, file: UploadFile, max_size_mb: int = 10) -> bool:
        """Validate that file size is within limits"""
        if hasattr(file, 'size') and file.size:
            return file.size <= (max_size_mb * 1024 * 1024)
        return True  # If we can't determine size, allow it